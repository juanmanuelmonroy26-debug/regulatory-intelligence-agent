from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pydantic import BaseModel

if TYPE_CHECKING:
    from src.models import InterpretedChange, SourceID


# ── Priority color mapping ─────────────────────────────────────────────────────

_PRIORITY_COLORS: dict[str, RGBColor] = {
    "Alta": RGBColor(0xC0, 0x00, 0x00),
    "Media": RGBColor(0xFF, 0x8C, 0x00),
    "Baja": RGBColor(0x00, 0x70, 0xC0),
}

_CHANGE_TYPE_LABELS: dict[str, str] = {
    "added": "Agregada",
    "removed": "Eliminada",
    "modified": "Modificada",
}

_SOURCE_LABELS: dict[str, str] = {
    "micrositios": "Micrositios DIAN — Normatividad",
    "normograma": "Normograma DIAN — Sistema de Facturación (ítem 1.6)",
    "proyectos_normas": "DIAN — Proyectos de Normas",
}


class ReportResult(BaseModel):
    report_path: str
    changes_included: int
    run_id: str
    generated_at: datetime


class DocumentReporter:
    def __init__(self, reports_dir: Path = Path("reports")) -> None:
        self.reports_dir = reports_dir
        self.reports_dir.mkdir(exist_ok=True)

    def generate(
        self,
        changes: list["InterpretedChange"],
        run_id: str,
        sources_monitored: list["SourceID"],
        no_changes_sources: list["SourceID"] | None = None,
        error_sources: dict[str, str] | None = None,
    ) -> ReportResult:
        doc = Document()
        generated_at = datetime.now(timezone.utc)

        self._set_document_styles(doc)
        self._write_header(doc, generated_at, run_id, sources_monitored)

        if changes:
            self._write_summary_table(doc, changes)
            for change in changes:
                self._write_change_section(doc, change)
        else:
            self._write_no_changes(doc, no_changes_sources or sources_monitored)

        if error_sources:
            self._write_errors(doc, error_sources)

        self._write_footer(doc, run_id, generated_at)

        date_str = generated_at.strftime("%Y-%m-%d")
        filename = f"reporte_dian_{date_str}_{run_id[:8]}.docx"
        report_path = self.reports_dir / filename
        doc.save(report_path)

        return ReportResult(
            report_path=str(report_path),
            changes_included=len(changes),
            run_id=run_id,
            generated_at=generated_at,
        )

    # ── Document-level helpers ─────────────────────────────────────────────────

    def _set_document_styles(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def _write_header(
        self,
        doc: Document,
        generated_at: datetime,
        run_id: str,
        sources: list["SourceID"],
    ) -> None:
        title = doc.add_heading("Reporte de Cambios Regulatorios DIAN", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_display = generated_at.strftime("%d de %B de %Y")
        subtitle = doc.add_paragraph(date_display)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

        doc.add_paragraph()

        meta = doc.add_paragraph()
        meta.add_run("Run ID: ").bold = True
        meta.add_run(run_id)
        meta.add_run("\nFuentes monitoreadas: ").bold = True
        meta.add_run(", ".join(_SOURCE_LABELS.get(s, s) for s in sources))

        doc.add_paragraph()

    def _write_summary_table(self, doc: Document, changes: list["InterpretedChange"]) -> None:
        doc.add_heading("Resumen", level=1)

        counts: dict[str, int] = {"Alta": 0, "Media": 0, "Baja": 0}
        for c in changes:
            counts[c.urgency] = counts.get(c.urgency, 0) + 1

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Total de cambios", "Prioridad Alta", "Prioridad Media", "Prioridad Baja"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].bold = True

        row = table.add_row().cells
        row[0].text = str(len(changes))
        row[1].text = str(counts["Alta"])
        row[2].text = str(counts["Media"])
        row[3].text = str(counts["Baja"])

        doc.add_paragraph()

    def _write_change_section(self, doc: Document, change: "InterpretedChange") -> None:
        diff = change.classified.diff
        change_label = _CHANGE_TYPE_LABELS.get(diff.change_type, diff.change_type)
        priority_label = change.urgency.value
        color = _PRIORITY_COLORS.get(priority_label, RGBColor(0, 0, 0))

        heading = doc.add_heading(level=2)
        heading.clear()
        run = heading.add_run(f"{diff.norm_id}  [{change_label}]")
        run.font.color.rgb = color

        # Metadata table
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = "Table Grid"
        rows_data = [
            ("Categoría funcional", change.classified.category.value),
            ("Módulo afectado", change.affected_module),
            ("Urgencia", priority_label),
            ("Equipos impactados", ", ".join(t.value for t in change.classified.affected_teams)),
        ]
        for i, (label, value) in enumerate(rows_data):
            cells = meta_table.rows[i].cells
            cells[0].text = label
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].text = str(value)
            if label == "Urgencia":
                cells[1].paragraphs[0].runs[0].font.color.rgb = color
                cells[1].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Narrative sections
        sections = [
            ("¿Qué cambió?", change.what_changed),
            ("Implicaciones", change.implications),
            ("Acción — Producto", change.product_action),
            ("Acción — Ingeniería", change.engineering_action),
            ("Acción — Customer Success", change.cs_action),
        ]
        for label, content in sections:
            p = doc.add_paragraph()
            p.add_run(label + ": ").bold = True
            p.add_run(content)

        doc.add_paragraph().add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def _write_no_changes(self, doc: Document, sources: list["SourceID"]) -> None:
        doc.add_heading("Sin cambios detectados", level=1)
        p = doc.add_paragraph(
            "No se detectaron modificaciones en las fuentes monitoreadas en esta ejecución."
        )
        p.add_run(
            f"\n\nFuentes revisadas: {', '.join(_SOURCE_LABELS.get(s, s) for s in sources)}"
        )

    def _write_errors(self, doc: Document, errors: dict[str, str]) -> None:
        doc.add_heading("Errores en esta ejecución", level=1)
        for source, msg in errors.items():
            p = doc.add_paragraph()
            p.add_run(f"{_SOURCE_LABELS.get(source, source)}: ").bold = True
            p.add_run(msg)

    def _write_footer(self, doc: Document, run_id: str, generated_at: datetime) -> None:
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Generado por Regulatory Intelligence Agent · "
            f"{generated_at.strftime('%Y-%m-%dT%H:%M:%SZ')} · run {run_id}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
